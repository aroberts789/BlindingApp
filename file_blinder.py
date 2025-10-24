#!/usr/bin/env python3
"""
File Blinding Script - Removes images and replaces keywords while preserving ALL structure
Supports: DOCX, HTML, TXT files
Keeps original structure, fonts, tables, styles - only removes images and replaces keywords
"""

import re
import os
import sys
from pathlib import Path
import zipfile
import tempfile
from xml.etree import ElementTree as ET
import hashlib

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup

    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False


class FileBlinder:
    def __init__(self, keyword_replacements=None, image_hashes_to_remove=None, standardize_formatting=True,
                 font_name="Calibri", font_size=11, font_color_black=True, grey_shading=False):
        """
        Initialize with keyword replacement dictionary and formatting options

        Args:
            keyword_replacements (dict): Dictionary of {original: replacement}
            standardize_formatting (bool): Whether to standardize fonts and colors
            font_name (str): Font to standardize to
            font_size (int): Font size to use (None to keep original)
            font_color_black (bool): Whether to make all text black
            grey_shading (bool): Whether to add grey shading to text
        """
        self.image_hashes_to_remove = set(image_hashes_to_remove or [])
        self.keyword_replacements = keyword_replacements or {
            "confidential": "[REDACTED]",
            "secret": "[REDACTED]",
            "internal": "[INTERNAL]",
            "proprietary": "[PROPRIETARY]",
            "classified": "[CLASSIFIED]",
            # Regex patterns
            r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b': '[EMAIL]',
            r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b': '[PHONE]',
            r'\b\d{3}-\d{2}-\d{4}\b': '[SSN]',
        }

        self.standardize_formatting = standardize_formatting
        self.font_name = font_name
        self.font_size = font_size
        self.font_color_black = font_color_black
        self.grey_shading = grey_shading

    def calculate_image_hash(self, image_data):
        """Calculate SHA256 hash of image data"""
        import hashlib
        return hashlib.sha256(image_data).hexdigest()

    def should_remove_image(self, image_data):
        """Check if an image should be removed based on its hash"""
        if not self.image_hashes_to_remove:
            return True  # Backward compatibility: remove all if no selection

        image_hash = self.calculate_image_hash(image_data)
        return image_hash in self.image_hashes_to_remove
    def extract_document_structure(self, input_path):
        """Extract document content with metadata for diff generation"""
        extension = Path(input_path).suffix.lower()

        if extension == '.docx':
            return self._extract_docx_structure(input_path)
        elif extension in ['.html', '.htm']:
            return self._extract_html_structure(input_path)
        elif extension == '.txt':
            return self._extract_txt_structure(input_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")

    def _extract_docx_structure(self, input_path):
        """Extract structure from DOCX file"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed")

        doc = Document(input_path)
        structure = {
            'type': 'docx',
            'paragraphs': [],
            'images': [],
            'tables': []
        }

        # Extract paragraphs with metadata
        for idx, para in enumerate(doc.paragraphs):
            para_data = {
                'index': idx,
                'text': para.text,
                'has_image': self._has_drawing_elements(para),
                'formatting': self._extract_paragraph_formatting(para),
                'style': para.style.name if para.style else 'Normal'
            }
            structure['paragraphs'].append(para_data)

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = {
                'index': table_idx,
                'rows': [],
                'has_shading': False
            }

            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text
                    has_shading = self._cell_has_shading(cell)
                    if has_shading:
                        table_data['has_shading'] = True
                    row_data.append({
                        'text': cell_text,
                        'has_shading': has_shading
                    })
                table_data['rows'].append(row_data)

            structure['tables'].append(table_data)

        return structure

    def _has_drawing_elements(self, paragraph):
        """Check if paragraph contains images"""
        try:
            p_element = paragraph._element
            for child in p_element.iter():
                if 'drawing' in str(child.tag).lower() or 'object' in str(child.tag).lower():
                    return True
            return False
        except:
            return False

    def _extract_paragraph_formatting(self, paragraph):
        """Extract formatting information from paragraph"""
        formatting = {
            'font_name': None,
            'font_size': None,
            'font_color': None,
            'is_bold': False,
            'is_italic': False
        }

        try:
            if paragraph.runs:
                first_run = paragraph.runs[0]
                if first_run.font.name:
                    formatting['font_name'] = first_run.font.name
                if first_run.font.size:
                    formatting['font_size'] = first_run.font.size.pt
                if first_run.font.color.rgb:
                    rgb = first_run.font.color.rgb
                    formatting['font_color'] = f"#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}"
                formatting['is_bold'] = first_run.font.bold or False
                formatting['is_italic'] = first_run.font.italic or False
        except:
            pass

        return formatting

    def _cell_has_shading(self, cell):
        """Check if table cell has background shading"""
        try:
            tc_element = cell._tc
            for child in tc_element.iter():
                tag_str = str(child.tag).lower()
                if 'shd' in tag_str:
                    return True
            return False
        except:
            return False

    def _extract_html_structure(self, input_path):
        """Extract structure from HTML file"""
        if not HTML_AVAILABLE:
            raise ImportError("beautifulsoup4 not installed")

        with open(input_path, 'r', encoding='utf-8') as file:
            content = file.read()

        soup = BeautifulSoup(content, 'html.parser')

        structure = {
            'type': 'html',
            'paragraphs': [],
            'images': []
        }

        # Extract text content
        for idx, para in enumerate(soup.find_all(['p', 'div', 'h1', 'h2', 'h3'])):
            structure['paragraphs'].append({
                'index': idx,
                'text': para.get_text(),
                'tag': para.name
            })

        # Count images
        images = soup.find_all(['img', 'picture', 'svg'])
        structure['images'] = [{'index': i} for i in range(len(images))]

        return structure

    def _extract_txt_structure(self, input_path):
        """Extract structure from TXT file"""
        try:
            with open(input_path, 'r', encoding='utf-8') as file:
                content = file.read()
        except UnicodeDecodeError:
            with open(input_path, 'r', encoding='latin-1') as file:
                content = file.read()

        paragraphs = content.split('\n\n')

        structure = {
            'type': 'txt',
            'paragraphs': [
                {'index': idx, 'text': para.strip()}
                for idx, para in enumerate(paragraphs) if para.strip()
            ]
        }

        return structure

    def generate_diff(self, original_structure, processed_structure):
        """Generate diff between original and processed structures"""
        import difflib

        diff_data = {
            'paragraph_changes': [],
            'image_changes': [],
            'table_changes': [],
            'formatting_changes': []
        }

        # Compare paragraphs
        for idx, (orig_para, proc_para) in enumerate(zip(
                original_structure.get('paragraphs', []),
                processed_structure.get('paragraphs', [])
        )):
            orig_text = orig_para.get('text', '')
            proc_text = proc_para.get('text', '')

            if orig_text != proc_text:
                # Generate character-level diff
                matcher = difflib.SequenceMatcher(None, orig_text, proc_text)
                text_changes = []

                for tag, i1, i2, j1, j2 in matcher.get_opcodes():
                    if tag == 'replace':
                        text_changes.append({
                            'type': 'replace',
                            'original': orig_text[i1:i2],
                            'processed': proc_text[j1:j2],
                            'position': i1
                        })
                    elif tag == 'delete':
                        text_changes.append({
                            'type': 'delete',
                            'original': orig_text[i1:i2],
                            'position': i1
                        })
                    elif tag == 'insert':
                        text_changes.append({
                            'type': 'insert',
                            'processed': proc_text[j1:j2],
                            'position': j1
                        })

                if text_changes:
                    diff_data['paragraph_changes'].append({
                        'index': idx,
                        'changes': text_changes
                    })

            # Check for image changes
            if orig_para.get('has_image') and not proc_para.get('has_image'):
                diff_data['image_changes'].append({
                    'paragraph_index': idx,
                    'removed': True
                })

            # Check for formatting changes
            orig_format = orig_para.get('formatting', {})
            proc_format = proc_para.get('formatting', {})

            if orig_format != proc_format:
                diff_data['formatting_changes'].append({
                    'paragraph_index': idx,
                    'original': orig_format,
                    'processed': proc_format
                })

        # Compare tables
        for idx, (orig_table, proc_table) in enumerate(zip(
                original_structure.get('tables', []),
                processed_structure.get('tables', [])
        )):
            if orig_table.get('has_shading') != proc_table.get('has_shading'):
                diff_data['table_changes'].append({
                    'index': idx,
                    'shading_removed': True
                })

        return diff_data

    def replace_keywords_in_text(self, text):
        """Replace keywords in text based on replacement dictionary"""
        if not text:
            return text

        for original, replacement in self.keyword_replacements.items():
            # Use regex for pattern matching
            if original.startswith(r'\b') or original.startswith(r'['):
                text = re.sub(original, replacement, text, flags=re.IGNORECASE)
            else:
                # Simple string replacement (case insensitive)
                text = re.sub(re.escape(original), replacement, text, flags=re.IGNORECASE)

        return text

    def remove_document_themes(self, doc):
        """Remove document themes that might cause colored text - AGGRESSIVE VERSION"""
        try:
            from docx.shared import RGBColor

            # Clear theme colors by setting document to a basic theme
            if hasattr(doc, 'settings'):
                try:
                    doc.settings.theme = None
                except:
                    pass

            # Also try to clear theme at the document level
            if hasattr(doc, '_element'):
                doc_element = doc._element

                # Find and remove theme elements
                themes_to_remove = []
                for child in doc_element.iter():
                    tag_str = str(child.tag).lower()
                    if any(theme_word in tag_str for theme_word in ['theme', 'themefont', 'clrscheme', 'fontscheme']):
                        themes_to_remove.append(child)

                for theme in themes_to_remove:
                    parent = theme.getparent()
                    if parent is not None:
                        try:
                            parent.remove(theme)
                        except:
                            pass

                # Remove any theme color references throughout the document
                namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

                # Find and remove theme color references (w:themeColor)
                try:
                    for color_elem in doc_element.xpath('.//w:color[@w:themeColor]', namespaces=namespaces):
                        # Remove the themeColor attribute
                        theme_color_attr = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeColor'
                        if theme_color_attr in color_elem.attrib:
                            del color_elem.attrib[theme_color_attr]
                        # Set explicit black color
                        color_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '000000')
                except:
                    pass

                # Remove theme fill references
                try:
                    for fill_elem in doc_element.xpath('.//w:shd[@w:themeFill]', namespaces=namespaces):
                        theme_fill_attr = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeFill'
                        if theme_fill_attr in fill_elem.attrib:
                            del fill_elem.attrib[theme_fill_attr]
                        # Remove the entire shading element
                        parent = fill_elem.getparent()
                        if parent is not None:
                            parent.remove(fill_elem)
                except:
                    pass

                # Remove theme tint/shade attributes
                try:
                    for elem in doc_element.xpath('.//*[@w:themeTint or @w:themeShade]', namespaces=namespaces):
                        for attr in ['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeTint',
                                     '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeShade']:
                            if attr in elem.attrib:
                                del elem.attrib[attr]
                except:
                    pass

        except Exception as e:
            # Continue if theme removal fails
            pass

    def standardize_run_formatting(self, run):
        """Apply standard formatting to a text run"""
        if not self.standardize_formatting:
            return

        try:
            # Set font name
            if self.font_name:
                run.font.name = self.font_name

            # Set font size if specified
            if self.font_size:
                from docx.shared import Pt
                run.font.size = Pt(self.font_size)

            # Aggressively set font color to black (remove all colors including theme colors)
            if self.font_color_black:
                from docx.shared import RGBColor
                run.font.color.rgb = RGBColor(0, 0, 0)  # Black

                # Also clear theme color more aggressively
                try:
                    run.font.color.theme_color = None
                    # Try to clear any color index
                    if hasattr(run.font.color, '_color_val'):
                        run.font.color._color_val = None
                except:
                    pass

            # Remove all highlighting and shading
            run.font.highlight_color = None  # Remove highlight

            # Also try to clear any theme-based highlighting
            try:
                if hasattr(run._element, 'rPr'):
                    rPr = run._element.rPr
                    if rPr is not None:
                        # Remove highlight elements
                        highlights_to_remove = []
                        for child in rPr:
                            if 'highlight' in str(child.tag).lower() or 'shd' in str(child.tag).lower():
                                highlights_to_remove.append(child)

                        for highlight in highlights_to_remove:
                            rPr.remove(highlight)
            except:
                pass

            # Remove underlines and other special formatting while keeping bold/italic
            # run.font.underline = None  # Uncomment if you want to remove underlines too

        except Exception as e:
            # If formatting fails, continue - text replacement is more important
            pass

    def remove_table_cell_shading(self, cell):
        """Remove background shading from a table cell - AGGRESSIVE VERSION"""
        try:
            # Work at the XML level to remove cell shading
            tc_element = cell._tc

            # Find and remove ALL shading/fill/background elements from the cell
            shading_elements_to_remove = []
            for child in tc_element.iter():
                tag_str = str(child.tag).lower()
                # Check for any shading-related tags
                if any(shading_word in tag_str for shading_word in ['shd', 'fill', 'background', 'bgcolor']):
                    shading_elements_to_remove.append(child)

            for shading_elem in shading_elements_to_remove:
                parent = shading_elem.getparent()
                if parent is not None:
                    try:
                        parent.remove(shading_elem)
                    except:
                        pass

            # Also check for table cell properties and remove shading using XPath
            try:
                namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

                # Remove shading from cell properties
                for tcPr in tc_element.xpath('.//w:tcPr', namespaces=namespaces):
                    # Remove any shading within table cell properties
                    for shd in tcPr.xpath('.//w:shd', namespaces=namespaces):
                        tcPr.remove(shd)

                    # Also remove any fill elements
                    for fill_elem in tcPr.xpath('.//*[contains(local-name(), "fill")]'):
                        try:
                            fill_elem.getparent().remove(fill_elem)
                        except:
                            pass
            except:
                pass

            # Additional cleanup: clear any attributes that might contain color
            try:
                if hasattr(tc_element, 'attrib'):
                    # Remove any color/fill attributes
                    attrs_to_remove = [k for k in tc_element.attrib.keys()
                                       if any(x in str(k).lower() for x in ['color', 'fill', 'shd', 'background'])]
                    for attr in attrs_to_remove:
                        del tc_element.attrib[attr]
            except:
                pass

        except Exception as e:
            # Continue if cell shading removal fails
            pass

    def remove_table_row_shading(self, row):
        """Remove background shading from a table row"""
        try:
            # Work at the XML level to remove row shading
            tr_element = row._tr

            # Find and remove ALL shading/fill/background elements from the row
            shading_elements_to_remove = []
            for child in tr_element.iter():
                tag_str = str(child.tag).lower()
                # Check for any shading-related tags
                if any(shading_word in tag_str for shading_word in ['shd', 'fill', 'background', 'bgcolor']):
                    shading_elements_to_remove.append(child)

            for shading_elem in shading_elements_to_remove:
                parent = shading_elem.getparent()
                if parent is not None:
                    try:
                        parent.remove(shading_elem)
                    except:
                        pass

            # Also check for table row properties and remove shading using XPath
            try:
                namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

                # Remove shading from row properties
                for trPr in tr_element.xpath('.//w:trPr', namespaces=namespaces):
                    # Remove any shading within table row properties
                    for shd in trPr.xpath('.//w:shd', namespaces=namespaces):
                        trPr.remove(shd)
            except:
                pass

        except Exception as e:
            # Continue if row shading removal fails
            pass

    def remove_content_control_shading(self, doc):
        """Remove background colors and styling from content controls - SURGICAL APPROACH"""
        try:
            from docx.shared import RGBColor
            from xml.etree import ElementTree as ET

            # Get the document element
            doc_element = doc._element

            namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                          'w15': 'http://schemas.microsoft.com/office/word/2012/wordml'}

            print("  Searching for content controls...")

            # Find all SDT (structured document tag) elements
            sdt_elements = []
            try:
                sdt_elements.extend(doc_element.xpath('.//w:sdt', namespaces=namespaces))
            except:
                pass

            for elem in doc_element.iter():
                if 'sdt' in str(elem.tag).lower():
                    if elem not in sdt_elements:
                        sdt_elements.append(elem)

            print(f"  Found {len(sdt_elements)} content controls")

            # Find all SDT properties
            sdtPr_elements = []
            try:
                sdtPr_elements.extend(doc_element.xpath('.//w:sdtPr', namespaces=namespaces))
            except:
                pass

            for elem in doc_element.iter():
                if 'sdtpr' in str(elem.tag).lower():
                    if elem not in sdtPr_elements:
                        sdtPr_elements.append(elem)

            print(f"  Found {len(sdtPr_elements)} SDT property elements")

            # Collect all unique SDT properties
            all_sdtPr_to_process = set()
            for sdt in sdt_elements:
                try:
                    for sdtPr in sdt.xpath('.//w:sdtPr', namespaces=namespaces):
                        all_sdtPr_to_process.add(sdtPr)
                except:
                    pass
                for child in sdt:
                    if 'sdtpr' in str(child.tag).lower():
                        all_sdtPr_to_process.add(child)

            for sdtPr in sdtPr_elements:
                all_sdtPr_to_process.add(sdtPr)

            print(f"  Processing {len(all_sdtPr_to_process)} unique SDT property elements...")

            # Process each SDT property - SURGICAL removal of only styling elements
            for sdtPr in all_sdtPr_to_process:
                try:
                    # List of element types to remove (these cause styling/borders)
                    elements_to_remove = []

                    for child in list(sdtPr):
                        tag_lower = str(child.tag).lower()
                        # Remove specific styling elements that cause borders/backgrounds
                        if any(x in tag_lower for x in [
                            'rpr',  # Run properties (text color, font)
                            'ppr',  # Paragraph properties
                            'color',  # Color elements
                            'shd',  # Shading
                            'fill',  # Fill color
                            'background',  # Background
                            'bdr',  # Border
                            'border'  # Border variant
                        ]):
                            elements_to_remove.append(child)
                            print(f"      Marking for removal: {child.tag}")

                    # Remove the styling elements
                    for elem in elements_to_remove:
                        try:
                            sdtPr.remove(elem)
                            print(f"      Removed: {elem.tag}")
                        except Exception as e:
                            print(f"      Could not remove {elem.tag}: {e}")

                    # Check if appearance element exists
                    appearance_exists = False
                    for child in sdtPr:
                        if 'appearance' in str(child.tag).lower():
                            # Update existing appearance to hidden
                            child.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'hidden')
                            appearance_exists = True
                            print(f"      Updated appearance to hidden")
                            break

                    # Add appearance="hidden" if it doesn't exist
                    if not appearance_exists:
                        appearance_elem = ET.Element(
                            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}appearance')
                        appearance_elem.set(
                            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'hidden')
                        sdtPr.insert(0, appearance_elem)
                        print(f"      Added appearance=hidden")

                    # Check if showingPlcHdr exists
                    showing_exists = False
                    for child in sdtPr:
                        if 'showingplchdr' in str(child.tag).lower():
                            # Update to not show placeholder
                            child.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0')
                            showing_exists = True
                            print(f"      Updated showingPlcHdr to 0")
                            break

                    # Add showingPlcHdr="0" if it doesn't exist
                    if not showing_exists:
                        showing_elem = ET.Element(
                            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}showingPlcHdr')
                        showing_elem.set(
                            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0')
                        # Insert after appearance if it exists
                        insert_pos = 1 if appearance_exists else 0
                        sdtPr.insert(insert_pos, showing_elem)
                        print(f"      Added showingPlcHdr=0")

                except Exception as e:
                    print(f"    Error processing SDT property: {e}")
                    import traceback
                    print(traceback.format_exc())

            # Process content inside SDTs
            print("  Removing styles from content inside all SDTs...")
            for sdt in sdt_elements:
                try:
                    for sdtContent in sdt.iter():
                        if 'sdtcontent' in str(sdtContent.tag).lower():
                            for para in sdtContent.iter():
                                if para.tag.endswith('}p') or 'p' == str(para.tag).split('}')[-1]:
                                    for pPr in para.findall('.//w:pPr', namespaces=namespaces):
                                        # Remove paragraph style references
                                        for pStyle in pPr.findall('.//w:pStyle', namespaces=namespaces):
                                            pPr.remove(pStyle)
                                            print(f"      Removed paragraph style from SDT content")
                                        # Remove shading
                                        for shd in pPr.findall('.//w:shd', namespaces=namespaces):
                                            pPr.remove(shd)
                                            print(f"      Removed paragraph shading from SDT content")
                                        # Remove borders
                                        for pBdr in pPr.findall('.//w:pBdr', namespaces=namespaces):
                                            pPr.remove(pBdr)
                                            print(f"      Removed paragraph border from SDT content")

                                    # Process runs
                                    for run in para.findall('.//w:r', namespaces=namespaces):
                                        for rPr in run.findall('.//w:rPr', namespaces=namespaces):
                                            # Remove run style references
                                            for rStyle in rPr.findall('.//w:rStyle', namespaces=namespaces):
                                                rPr.remove(rStyle)
                                                print(f"      Removed run style from SDT content")
                                            # Remove shading
                                            for shd in rPr.findall('.//w:shd', namespaces=namespaces):
                                                rPr.remove(shd)
                                                print(f"      Removed run shading from SDT content")
                except Exception as e:
                    print(f"  Error processing SDT content: {e}")

        except Exception as e:
            print(f"  Content control shading removal error: {e}")
            import traceback
            print(traceback.format_exc())

    def remove_paragraph_borders(self, paragraph):
        """Remove paragraph borders"""
        try:
            # Remove borders at the paragraph format level
            paragraph_format = paragraph.paragraph_format

            # Clear paragraph borders
            if hasattr(paragraph_format, 'borders'):
                borders = paragraph_format.borders
                if borders:
                    try:
                        borders.top = None
                        borders.bottom = None
                        borders.left = None
                        borders.right = None
                    except:
                        pass

            # Also work at the XML level to remove border elements
            p_element = paragraph._element

            # Find and remove all border-related elements
            border_elements_to_remove = []
            for child in p_element.iter():
                tag_str = str(child.tag).lower()
                if any(border_word in tag_str for border_word in ['bdr', 'border', 'pBdr']):
                    border_elements_to_remove.append(child)

            for border_elem in border_elements_to_remove:
                parent = border_elem.getparent()
                if parent is not None:
                    parent.remove(border_elem)

            # Also check for paragraph properties and remove borders
            try:
                for pPr in p_element.xpath('.//w:pPr', namespaces={
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    # Remove any borders within paragraph properties
                    for border in pPr.xpath('.//w:pBdr', namespaces={
                        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                        pPr.remove(border)
            except:
                pass

        except Exception as e:
            # Continue if border removal fails
            pass

    def remove_paragraph_shading(self, paragraph):
        """Remove paragraph-level shading and background colors"""
        try:
            # Remove paragraph shading at multiple levels
            paragraph_format = paragraph.paragraph_format

            # Clear paragraph-level shading
            if hasattr(paragraph_format, 'shading'):
                shading = paragraph_format.shading
                if shading:
                    shading.background_pattern_color = None
                    shading.foreground_pattern_color = None
                    # Also try to clear the fill
                    try:
                        shading.fill = None
                    except:
                        pass

            # Also work at the XML level to remove shading elements
            p_element = paragraph._element

            # Find and remove all shading-related elements
            shading_elements_to_remove = []
            for child in p_element.iter():
                tag_str = str(child.tag).lower()
                if any(shading_word in tag_str for shading_word in ['shd', 'fill', 'highlight', 'bgcolor']):
                    shading_elements_to_remove.append(child)

            for shading_elem in shading_elements_to_remove:
                parent = shading_elem.getparent()
                if parent is not None:
                    parent.remove(shading_elem)

            # Also check for paragraph properties and remove background colors
            try:
                for pPr in p_element.xpath('.//w:pPr', namespaces={
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    # Remove any shading within paragraph properties
                    for shd in pPr.xpath('.//w:shd', namespaces={
                        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                        pPr.remove(shd)
            except:
                pass

        except Exception as e:
            # Continue if shading removal fails
            pass

    def remove_hyperlinks_from_paragraph(self, paragraph):
        """Remove hyperlinks from a paragraph while preserving the text content"""
        try:
            from docx.shared import RGBColor

            p_element = paragraph._element

            # Find all hyperlink elements
            hyperlinks = []
            for child in p_element.iter():
                if 'hyperlink' in str(child.tag).lower():
                    hyperlinks.append(child)

            # Process each hyperlink
            for hyperlink in hyperlinks:
                # Extract all text runs from the hyperlink before removing it
                # Hyperlinks contain runs (w:r elements) that have the actual text
                parent = hyperlink.getparent()
                if parent is not None:
                    # Get the position of the hyperlink in the parent
                    hyperlink_index = list(parent).index(hyperlink)

                    # Extract all child elements (runs) from the hyperlink
                    children_to_preserve = list(hyperlink)

                    # Process each run to remove hyperlink formatting (underline, blue color)
                    for child in children_to_preserve:
                        # Look for run properties (rPr) within each run
                        try:
                            for rPr in child.iter():
                                if 'rPr' in str(rPr.tag):
                                    # Remove underline elements
                                    underlines_to_remove = []
                                    for elem in rPr:
                                        if 'u' in str(elem.tag).lower() and 'u' == str(elem.tag).split('}')[-1]:
                                            underlines_to_remove.append(elem)

                                    for u_elem in underlines_to_remove:
                                        rPr.remove(u_elem)

                                    # Remove color elements (the blue hyperlink color)
                                    colors_to_remove = []
                                    for elem in rPr:
                                        if 'color' in str(elem.tag).lower():
                                            colors_to_remove.append(elem)

                                    for color_elem in colors_to_remove:
                                        rPr.remove(color_elem)
                        except:
                            pass

                    # Insert the runs directly into the paragraph where the hyperlink was
                    for i, child in enumerate(children_to_preserve):
                        parent.insert(hyperlink_index + i, child)

                    # Now remove the empty hyperlink element
                    parent.remove(hyperlink)

            # After removing hyperlinks, process all runs again to ensure formatting
            for run in paragraph.runs:
                try:
                    # Remove underline
                    run.font.underline = None

                    # Set color to black
                    if self.font_color_black:
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        try:
                            run.font.color.theme_color = None
                        except:
                            pass
                except:
                    pass

        except Exception as e:
            # Continue if hyperlink removal fails
            pass

    def remove_list_formatting(self, paragraph):
        """Remove list bullet highlighting and formatting"""
        try:
            # Clear list formatting - reset to normal paragraph
            paragraph.style = paragraph._document.styles['Normal']

            # Remove paragraph numbering properties
            paragraph_format = paragraph.paragraph_format

            # Clear left indent that might cause bullet appearance
            paragraph_format.left_indent = None
            paragraph_format.first_line_indent = None

            # Remove any numbering
            p_element = paragraph._element

            # Find and remove numbering properties
            numbering_to_remove = []
            for child in p_element.iter():
                if 'numPr' in str(child.tag) or 'pPr' in str(child.tag):
                    # Look for numbering inside pPr
                    for subchild in child:
                        if 'numPr' in str(subchild.tag):
                            numbering_to_remove.append(subchild)

            for num_element in numbering_to_remove:
                parent = num_element.getparent()
                if parent is not None:
                    parent.remove(num_element)

        except Exception as e:
            # Continue if list formatting removal fails
            pass

    def process_docx_safe(self, input_path, output_path):
        """Enhanced DOCX processing with aggressive content control and color removal"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        print("=" * 70)
        print("ENHANCED SAFE MODE - Complete Content Control & Color Removal")
        print("=" * 70)
        print()

        # STEP 0: Remove selected images FIRST (simple, focused deletion)
        input_path = Path(input_path)
        temp_no_images = self.remove_selected_images_simple(input_path)

        # Now continue with the rest of processing using the image-free version
        print("\nContinuing with formatting cleanup...\n")

        # STEP 1: PRE-PROCESS AT XML LEVEL (Remove content controls, etc.)
        print("Step 1: Pre-processing at XML level...")
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_dir = Path(temp_dir)

            # Extract DOCX
            print("  Extracting DOCX...")
            with zipfile.ZipFile(temp_no_images, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)

            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            }

            for prefix, uri in namespaces.items():
                ET.register_namespace(prefix, uri)

            # Process document.xml to remove content controls
            document_xml = temp_dir / 'word' / 'document.xml'
            sdts_removed = 0

            if document_xml.exists():
                print("  Removing content controls from document.xml...")
                tree = ET.parse(document_xml)
                root = tree.getroot()
                parent_map = {c: p for p in tree.iter() for c in p}

                # Find and remove ALL content controls
                for sdt in root.findall('.//w:sdt', namespaces):
                    parent = parent_map.get(sdt)
                    if parent is not None:
                        sdt_index = list(parent).index(sdt)

                        # Extract content from SDT (preserves tables and everything)
                        sdt_content = sdt.find('.//w:sdtContent', namespaces)
                        if sdt_content is not None:
                            # Move all children from sdtContent to parent
                            for child in list(sdt_content):
                                parent.insert(sdt_index, child)
                                sdt_index += 1

                        # Remove the SDT wrapper entirely
                        parent.remove(sdt)
                        sdts_removed += 1

                # REMOVE ALL PARAGRAPH STYLES (force everything to Normal)
                print("  Removing all paragraph styles (forcing to Normal)...")
                styles_removed = 0
                for para in root.findall('.//w:p', namespaces):
                    for pPr in para.findall('.//w:pPr', namespaces):
                        # Remove paragraph style references
                        for pStyle in pPr.findall('.//w:pStyle', namespaces):
                            pPr.remove(pStyle)
                            styles_removed += 1

                # Force ALL text to black color
                print("  Forcing all text to black...")
                colors_forced = 0
                for rPr in root.findall('.//w:rPr', namespaces):
                    # Remove existing color elements
                    for color_elem in list(rPr):
                        if 'color' in str(color_elem.tag).lower() or 'highlight' in str(color_elem.tag).lower():
                            rPr.remove(color_elem)

                    # Add black color
                    color_elem = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
                    color_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '000000')

                    # Remove theme color attributes
                    for attr in ['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeColor',
                                 '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeTint',
                                 '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeShade']:
                        if attr in color_elem.attrib:
                            del color_elem.attrib[attr]

                    rPr.insert(0, color_elem)
                    colors_forced += 1

                # Remove ALL shading
                print("  Removing all shading...")
                shading_removed = 0
                for shd in root.findall('.//w:shd', namespaces):
                    parent = parent_map.get(shd)
                    if parent is not None:
                        parent.remove(shd)
                        shading_removed += 1

                # Remove ALL paragraph borders
                print("  Removing all paragraph borders...")
                borders_removed = 0
                for pBdr in root.findall('.//w:pBdr', namespaces):
                    parent = parent_map.get(pBdr)
                    if parent is not None:
                        parent.remove(pBdr)
                        borders_removed += 1

                # Save modified document.xml
                tree.write(document_xml, encoding='utf-8', xml_declaration=True)

                print(f"  ✓ Removed {sdts_removed} content controls")
                print(f"  ✓ Removed {styles_removed} paragraph styles (forced to Normal)")
                print(f"  ✓ Forced {colors_forced} text colors to black")
                print(f"  ✓ Removed {shading_removed} shading elements")
                print(f"  ✓ Removed {borders_removed} border elements")

            # Neutralize theme files (don't delete - python-docx expects them)
            print("  Neutralizing theme files...")
            theme_dir = temp_dir / 'word' / 'theme'
            if theme_dir.exists():
                for theme_file in theme_dir.glob('*.xml'):
                    try:
                        tree = ET.parse(theme_file)
                        root = tree.getroot()

                        # Replace all colors with black
                        for element in root.iter():
                            if 'srgbClr' in str(element.tag):
                                element.set('val', '000000')
                            elif 'sysClr' in str(element.tag):
                                element.set('val', 'windowText')
                                element.set('lastClr', '000000')

                        tree.write(theme_file, encoding='utf-8', xml_declaration=True)
                    except:
                        pass
                print("  ✓ Neutralized theme files (colors set to black)")

            # Neutralize styles.xml
            print("  Neutralizing styles.xml...")
            styles_xml = temp_dir / 'word' / 'styles.xml'
            if styles_xml.exists():
                tree = ET.parse(styles_xml)
                root = tree.getroot()
                parent_map = {c: p for p in tree.iter() for c in p}

                for element in root.iter():
                    if 'color' in str(element.tag).lower() or 'shd' in str(element.tag).lower():
                        parent = parent_map.get(element)
                        if parent is not None:
                            try:
                                parent.remove(element)
                            except:
                                pass

                tree.write(styles_xml, encoding='utf-8', xml_declaration=True)
                print("  ✓ Neutralized styles.xml")

            # Rebuild DOCX to temp file
            print("  Rebuilding DOCX...")
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_preprocessed:
                preprocessed_path = tmp_preprocessed.name

            with zipfile.ZipFile(preprocessed_path, 'w', zipfile.ZIP_DEFLATED) as zip_out:
                for file_path in temp_dir.rglob('*'):
                    if file_path.is_file():
                        arc_path = file_path.relative_to(temp_dir)
                        zip_out.write(file_path, arc_path)

            print("  ✓ Pre-processing complete\n")

        # Clean up the temp file from step 0
        try:
            if temp_no_images != input_path:  # Don't delete original
                os.unlink(temp_no_images)
        except:
            pass

        # STEP 2: PROCESS WITH PYTHON-DOCX (for remaining cleanup)
        print("Step 2: Processing with python-docx...")
        print("Loading document...")
        doc = Document(preprocessed_path)

        text_replacements = 0
        styles_reset = 0

        # Process paragraphs
        print("Processing paragraphs...")
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if para_idx % 10 == 0:
                print(f"  Processing paragraph {para_idx + 1}/{len(doc.paragraphs)}")

            # FORCE PARAGRAPH STYLE TO NORMAL (removes heading styles, etc.)
            try:
                if paragraph.style.name != 'Normal':
                    paragraph.style = doc.styles['Normal']
                    styles_reset += 1
            except:
                pass

            # Process text in runs
            for run in paragraph.runs:
                if run.text:
                    original_text = run.text
                    new_text = self.replace_keywords_in_text(original_text)
                    if new_text != original_text:
                        run.text = new_text
                        text_replacements += 1

                # Apply formatting standardization
                self.standardize_run_formatting(run)

            # Remove hyperlinks, list formatting, borders, shading
            self.remove_hyperlinks_from_paragraph(paragraph)
            self.remove_list_formatting(paragraph)
            self.remove_paragraph_borders(paragraph)
            self.remove_paragraph_shading(paragraph)

        # Process tables
        print("Processing tables...")
        for table_idx, table in enumerate(doc.tables):
            print(f"  Processing table {table_idx + 1}/{len(doc.tables)}")
            for row in table.rows:
                self.remove_table_row_shading(row)

                for cell in row.cells:
                    self.remove_table_cell_shading(cell)

                    for paragraph in cell.paragraphs:
                        # FORCE PARAGRAPH STYLE TO NORMAL in tables too
                        try:
                            if paragraph.style.name != 'Normal':
                                paragraph.style = doc.styles['Normal']
                                styles_reset += 1
                        except:
                            pass

                        # Process text in cell runs
                        for run in paragraph.runs:
                            if run.text:
                                original_text = run.text
                                new_text = self.replace_keywords_in_text(original_text)
                                if new_text != original_text:
                                    run.text = new_text
                                    text_replacements += 1

                            self.standardize_run_formatting(run)

                        self.remove_hyperlinks_from_paragraph(paragraph)
                        self.remove_list_formatting(paragraph)
                        self.remove_paragraph_borders(paragraph)
                        self.remove_paragraph_shading(paragraph)

        # Process headers and footers
        print("Processing headers and footers...")
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    # FORCE PARAGRAPH STYLE TO NORMAL in headers
                    try:
                        if paragraph.style.name != 'Normal':
                            paragraph.style = doc.styles['Normal']
                            styles_reset += 1
                    except:
                        pass

                    for run in paragraph.runs:
                        if run.text:
                            original_text = run.text
                            new_text = self.replace_keywords_in_text(original_text)
                            if new_text != original_text:
                                run.text = new_text
                                text_replacements += 1

                        self.standardize_run_formatting(run)

                    self.remove_hyperlinks_from_paragraph(paragraph)

            if section.footer:
                for paragraph in section.footer.paragraphs:
                    # FORCE PARAGRAPH STYLE TO NORMAL in footers
                    try:
                        if paragraph.style.name != 'Normal':
                            paragraph.style = doc.styles['Normal']
                            styles_reset += 1
                    except:
                        pass

                    for run in paragraph.runs:
                        if run.text:
                            original_text = run.text
                            new_text = self.replace_keywords_in_text(original_text)
                            if new_text != original_text:
                                run.text = new_text
                                text_replacements += 1

                        self.standardize_run_formatting(run)

                    self.remove_hyperlinks_from_paragraph(paragraph)

        print("Saving final document...")
        doc.save(output_path)

        # Clean up temp file
        try:
            os.unlink(preprocessed_path)
        except:
            pass

        print(f"\n{'=' * 70}")
        print("✅ PROCESSING COMPLETE")
        print(f"{'=' * 70}")
        print(f"✓ Removed {sdts_removed} content controls (blue sections)")
        print(f"✓ Reset {styles_reset} paragraphs to Normal style (removed headings)")
        print(f"✓ Made {text_replacements} text replacements")
        print(f"✓ Forced all text to black")
        print(f"✓ Removed all shading and borders")
        print(f"✓ Tables preserved with structure intact")
        print()

        return output_path
    def process_docx_xml_safe(self, input_path, output_path):
        """Process DOCX by safely modifying XML while preserving structure"""

        # Create a temporary working directory
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_dir = Path(temp_dir)

            # Extract the DOCX file
            with zipfile.ZipFile(input_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)

            # Define XML namespace
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
            }

            # Register namespaces
            for prefix, uri in namespaces.items():
                ET.register_namespace(prefix, uri)

            images_removed = 0
            text_replacements = 0
            hyperlinks_removed = 0

            # NEUTRALIZE THEME FILES - Replace theme colors with black/white
            print("Neutralizing theme colors...")
            theme_dir = temp_dir / 'word' / 'theme'
            if theme_dir.exists():
                for theme_file in theme_dir.glob('*.xml'):
                    try:
                        tree = ET.parse(theme_file)
                        root = tree.getroot()

                        # Find all color scheme elements and replace with neutral colors
                        for color_scheme in root.iter():
                            tag = str(color_scheme.tag)
                            # Replace theme colors with black or white
                            if 'clrScheme' in tag or 'color' in tag.lower():
                                for color_elem in color_scheme:
                                    # Set all colors to either black (000000) or white (FFFFFF)
                                    for child in color_elem:
                                        if 'srgbClr' in str(child.tag):
                                            child.set('val', '000000')  # Black
                                        elif 'sysClr' in str(child.tag):
                                            child.set('val', 'windowText')
                                            child.set('lastClr', '000000')

                        tree.write(theme_file, encoding='utf-8', xml_declaration=True)
                        print(f"  Neutralized {theme_file.name}")
                    except Exception as e:
                        print(f"  Could not process theme file {theme_file.name}: {e}")

            # NEUTRALIZE STYLES.XML - Remove theme color references
            print("Neutralizing style theme references...")
            styles_xml = temp_dir / 'word' / 'styles.xml'
            if styles_xml.exists():
                try:
                    tree = ET.parse(styles_xml)
                    root = tree.getroot()

                    # Build a parent map since ElementTree doesn't have getparent()
                    parent_map = {c: p for p in tree.iter() for c in p}

                    # First pass: Remove all theme color attributes
                    for elem in root.iter():
                        # Remove theme color attributes
                        attrs_to_remove = []
                        for attr_name in list(elem.attrib.keys()):
                            if 'theme' in attr_name.lower() and 'color' in attr_name.lower():
                                attrs_to_remove.append(attr_name)

                        for attr in attrs_to_remove:
                            del elem.attrib[attr]

                    # Second pass: Mark shading elements for removal (don't remove while iterating)
                    elements_to_remove = []
                    for elem in root.iter():
                        if 'shd' in str(elem.tag).lower():
                            elements_to_remove.append(elem)

                    # Third pass: Remove marked elements using parent map
                    for elem in elements_to_remove:
                        parent = parent_map.get(elem)
                        if parent is not None:
                            try:
                                parent.remove(elem)
                            except:
                                pass

                    tree.write(styles_xml, encoding='utf-8', xml_declaration=True)
                    print(f"  Neutralized styles.xml")
                except Exception as e:
                    print(f"  Could not process styles.xml: {e}")
                    import traceback
                    traceback.print_exc()

            # Process main document
            document_xml = temp_dir / 'word' / 'document.xml'
            if document_xml.exists():
                print("Processing main document XML...")
                tree = ET.parse(document_xml)
                root = tree.getroot()

                # Build parent map for element removal
                parent_map = {c: p for p in tree.iter() for c in p}

                # Remove drawing elements (images)
                drawings_to_remove = root.findall('.//w:drawing', namespaces)
                for drawing in drawings_to_remove:
                    parent = parent_map.get(drawing)
                    if parent is not None:
                        parent.remove(drawing)
                        images_removed += 1

                # Remove object elements
                objects_to_remove = root.findall('.//w:object', namespaces)
                for obj in objects_to_remove:
                    parent = parent_map.get(obj)
                    if parent is not None:
                        parent.remove(obj)
                        images_removed += 1

                # Remove all shading elements (table cells, rows, paragraphs)
                shading_to_remove = root.findall('.//w:shd', namespaces)
                for shd in shading_to_remove:
                    parent = parent_map.get(shd)
                    if parent is not None:
                        try:
                            parent.remove(shd)
                        except:
                            pass

                # FORCE ALL TEXT TO BLACK COLOR
                print("Forcing all text to black color...")
                for rPr in root.findall('.//w:rPr', namespaces):
                    # Find or create color element
                    color_elem = rPr.find('w:color', namespaces)
                    if color_elem is None:
                        # Create new color element
                        color_elem = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
                        rPr.insert(0, color_elem)

                    # Set to black and remove theme color
                    color_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '000000')

                    # Remove theme color attributes if they exist
                    theme_color_attr = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeColor'
                    theme_tint_attr = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeTint'
                    theme_shade_attr = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeShade'

                    if theme_color_attr in color_elem.attrib:
                        del color_elem.attrib[theme_color_attr]
                    if theme_tint_attr in color_elem.attrib:
                        del color_elem.attrib[theme_tint_attr]
                    if theme_shade_attr in color_elem.attrib:
                        del color_elem.attrib[theme_shade_attr]

                # Remove content control (SDT) appearance/color properties AND BORDERS
                print("Removing content control styling...")
                for sdt in root.findall('.//w:sdt', namespaces):
                    try:
                        for sdtPr in sdt.findall('.//w:sdtPr', namespaces):
                            # Build parent map for this subtree
                            sdt_parent_map = {c: p for p in sdtPr.iter() for c in p}

                            # REMOVE STYLE REFERENCES - this is what causes the blue background!
                            print("  Resetting content control style...")
                            # Remove run properties (character styles)
                            for rPrElem in sdtPr.findall('.//w:rPr', namespaces):
                                parent = sdt_parent_map.get(rPrElem, sdtPr)
                                if parent is not None:
                                    try:
                                        parent.remove(rPrElem)
                                        print("    Removed rPr (run properties/style) from SDT")
                                    except:
                                        pass

                            # Remove paragraph properties (paragraph styles)
                            for pPrElem in sdtPr.findall('.//w:pPr', namespaces):
                                parent = sdt_parent_map.get(pPrElem, sdtPr)
                                if parent is not None:
                                    try:
                                        parent.remove(pPrElem)
                                        print("    Removed pPr (paragraph properties/style) from SDT")
                                    except:
                                        pass

                            # Remove any direct children that are style-related
                            direct_children_to_remove = []
                            for child in list(sdtPr):
                                tag_lower = str(child.tag).lower()
                                if 'rpr' in tag_lower or 'ppr' in tag_lower:
                                    direct_children_to_remove.append(child)
                                    print(f"    Marking style child for removal: {child.tag}")

                            for child in direct_children_to_remove:
                                try:
                                    sdtPr.remove(child)
                                    print(f"    Removed style child: {child.tag}")
                                except:
                                    pass

                            # SET appearance to hidden (removes border)
                            appearance_found = False
                            for appearance in sdtPr.findall('.//w:appearance', namespaces):
                                # Set appearance to "hidden" to remove border
                                appearance.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val',
                                               'hidden')
                                appearance_found = True
                                print("    Set appearance to hidden")

                            # If no appearance element exists, create one set to hidden
                            if not appearance_found:
                                appearance_elem = ET.Element(
                                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}appearance')
                                appearance_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val',
                                                    'hidden')
                                sdtPr.insert(0, appearance_elem)
                                print("    Created hidden appearance")

                            # Remove color elements
                            for color in sdtPr.findall('.//w:color', namespaces):
                                parent = sdt_parent_map.get(color, sdtPr)
                                if parent is not None:
                                    try:
                                        parent.remove(color)
                                    except:
                                        pass

                            # Remove any shading in SDT properties
                            for shd in sdtPr.findall('.//w:shd', namespaces):
                                parent = sdt_parent_map.get(shd, sdtPr)
                                if parent is not None:
                                    try:
                                        parent.remove(shd)
                                    except:
                                        pass

                            # Remove border-related elements more aggressively
                            all_children = list(sdtPr)
                            for child in all_children:
                                tag_str = str(child.tag).lower()
                                if any(x in tag_str for x in ['bdr', 'border']):
                                    try:
                                        sdtPr.remove(child)
                                    except:
                                        pass

                        # NOW ALSO PROCESS THE CONTENT INSIDE THE SDT (sdtContent)
                        # This is where the paragraph style that causes the blue background lives!
                        print("  Removing styles from content inside SDT...")
                        for sdtContent in sdt.findall('.//w:sdtContent', namespaces):
                            # Find all paragraphs inside the content
                            for para in sdtContent.findall('.//w:p', namespaces):
                                # Find paragraph properties
                                for pPr in para.findall('.//w:pPr', namespaces):
                                    # Remove paragraph style references (w:pStyle)
                                    for pStyle in pPr.findall('.//w:pStyle', namespaces):
                                        pPr.remove(pStyle)
                                        print(f"    Removed paragraph style reference from content")

                                    # Remove shading from paragraph
                                    for shd in pPr.findall('.//w:shd', namespaces):
                                        pPr.remove(shd)
                                        print(f"    Removed shading from paragraph")

                                # Also process runs inside these paragraphs
                                for run in para.findall('.//w:r', namespaces):
                                    for rPr in run.findall('.//w:rPr', namespaces):
                                        # Remove run style references (w:rStyle)
                                        for rStyle in rPr.findall('.//w:rStyle', namespaces):
                                            rPr.remove(rStyle)
                                            print(f"    Removed run style reference from content")

                                        # Remove shading from runs
                                        for shd in rPr.findall('.//w:shd', namespaces):
                                            rPr.remove(shd)
                                            print(f"    Removed shading from run")

                    except Exception as e:
                        print(f"  Error processing SDT: {e}")
                        import traceback
                        traceback.print_exc()

                # Remove hyperlinks while preserving text content
                for hyperlink in root.findall('.//w:hyperlink', namespaces):
                    parent = parent_map.get(hyperlink)
                    if parent is not None:
                        # Get the position of the hyperlink
                        hyperlink_index = list(parent).index(hyperlink)

                        # Move all children (runs) from hyperlink to parent
                        # AND remove hyperlink formatting (blue color, underline)
                        for i, child in enumerate(list(hyperlink)):
                            # If this is a run (w:r), clean up its formatting
                            if 'r' in str(child.tag).lower() and 'r' == str(child.tag).split('}')[-1]:
                                # Find run properties
                                for rPr in child.findall('w:rPr', namespaces):
                                    # Remove underline
                                    underline_elems = rPr.findall('w:u', namespaces)
                                    for u_elem in underline_elems:
                                        rPr.remove(u_elem)

                                    # Force color to black and remove theme color
                                    color_elem = rPr.find('w:color', namespaces)
                                    if color_elem is None:
                                        color_elem = ET.Element(
                                            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
                                        rPr.insert(0, color_elem)

                                    color_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val',
                                                   '000000')

                                    # Remove theme color attributes
                                    for attr in [
                                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeColor',
                                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeTint',
                                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeShade']:
                                        if attr in color_elem.attrib:
                                            del color_elem.attrib[attr]

                            parent.insert(hyperlink_index + i, child)

                        # Remove the hyperlink element
                        parent.remove(hyperlink)
                        hyperlinks_removed += 1

                # Replace text in text elements
                for text_elem in root.findall('.//w:t', namespaces):
                    if text_elem.text:
                        original_text = text_elem.text
                        new_text = self.replace_keywords_in_text(original_text)
                        if new_text != original_text:
                            text_elem.text = new_text
                            text_replacements += 1

                # Save the modified XML
                tree.write(document_xml, encoding='utf-8', xml_declaration=True)

            # Process headers
            header_files = list((temp_dir / 'word').glob('header*.xml'))
            for header_file in header_files:
                print(f"Processing {header_file.name}...")
                tree = ET.parse(header_file)
                root = tree.getroot()

                # Build parent map
                parent_map = {c: p for p in tree.iter() for c in p}

                # Remove images
                for drawing in root.findall('.//w:drawing', namespaces):
                    parent = parent_map.get(drawing)
                    if parent is not None:
                        parent.remove(drawing)
                        images_removed += 1

                # Remove all shading elements
                for shd in root.findall('.//w:shd', namespaces):
                    parent = parent_map.get(shd)
                    if parent is not None:
                        try:
                            parent.remove(shd)
                        except:
                            pass

                # FORCE ALL TEXT TO BLACK COLOR
                for rPr in root.findall('.//w:rPr', namespaces):
                    color_elem = rPr.find('w:color', namespaces)
                    if color_elem is None:
                        color_elem = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
                        rPr.insert(0, color_elem)
                    color_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '000000')
                    # Remove theme attributes
                    for attr in ['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeColor',
                                 '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeTint',
                                 '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeShade']:
                        if attr in color_elem.attrib:
                            del color_elem.attrib[attr]

                # Remove hyperlinks while preserving text
                for hyperlink in root.findall('.//w:hyperlink', namespaces):
                    parent = parent_map.get(hyperlink)
                    if parent is not None:
                        hyperlink_index = list(parent).index(hyperlink)
                        # Clean up formatting in runs from hyperlinks
                        for i, child in enumerate(list(hyperlink)):
                            if 'r' in str(child.tag).lower() and 'r' == str(child.tag).split('}')[-1]:
                                for rPr in child.findall('w:rPr', namespaces):
                                    # Remove underline
                                    for u_elem in rPr.findall('w:u', namespaces):
                                        rPr.remove(u_elem)
                                    # Force color to black
                                    color_elem = rPr.find('w:color', namespaces)
                                    if color_elem is None:
                                        color_elem = ET.Element(
                                            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
                                        rPr.insert(0, color_elem)
                                    color_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val',
                                                   '000000')
                                    for attr in [
                                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeColor',
                                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeTint',
                                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeShade']:
                                        if attr in color_elem.attrib:
                                            del color_elem.attrib[attr]
                            parent.insert(hyperlink_index + i, child)
                        parent.remove(hyperlink)
                        hyperlinks_removed += 1

                # Replace text
                for text_elem in root.findall('.//w:t', namespaces):
                    if text_elem.text:
                        original_text = text_elem.text
                        new_text = self.replace_keywords_in_text(original_text)
                        if new_text != original_text:
                            text_elem.text = new_text
                            text_replacements += 1

                tree.write(header_file, encoding='utf-8', xml_declaration=True)

            # Process footers
            footer_files = list((temp_dir / 'word').glob('footer*.xml'))
            for footer_file in footer_files:
                print(f"Processing {footer_file.name}...")
                tree = ET.parse(footer_file)
                root = tree.getroot()

                # Build parent map
                parent_map = {c: p for p in tree.iter() for c in p}

                # Remove images
                for drawing in root.findall('.//w:drawing', namespaces):
                    parent = parent_map.get(drawing)
                    if parent is not None:
                        parent.remove(drawing)
                        images_removed += 1

                # Remove all shading elements
                for shd in root.findall('.//w:shd', namespaces):
                    parent = parent_map.get(shd)
                    if parent is not None:
                        try:
                            parent.remove(shd)
                        except:
                            pass

                # FORCE ALL TEXT TO BLACK COLOR
                for rPr in root.findall('.//w:rPr', namespaces):
                    color_elem = rPr.find('w:color', namespaces)
                    if color_elem is None:
                        color_elem = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
                        rPr.insert(0, color_elem)
                    color_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '000000')
                    # Remove theme attributes
                    for attr in ['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeColor',
                                 '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeTint',
                                 '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeShade']:
                        if attr in color_elem.attrib:
                            del color_elem.attrib[attr]

                # Remove hyperlinks while preserving text
                for hyperlink in root.findall('.//w:hyperlink', namespaces):
                    parent = parent_map.get(hyperlink)
                    if parent is not None:
                        hyperlink_index = list(parent).index(hyperlink)
                        # Clean up formatting in runs from hyperlinks
                        for i, child in enumerate(list(hyperlink)):
                            if 'r' in str(child.tag).lower() and 'r' == str(child.tag).split('}')[-1]:
                                for rPr in child.findall('w:rPr', namespaces):
                                    # Remove underline
                                    for u_elem in rPr.findall('w:u', namespaces):
                                        rPr.remove(u_elem)
                                    # Force color to black
                                    color_elem = rPr.find('w:color', namespaces)
                                    if color_elem is None:
                                        color_elem = ET.Element(
                                            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
                                        rPr.insert(0, color_elem)
                                    color_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val',
                                                   '000000')
                                    for attr in [
                                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeColor',
                                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeTint',
                                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeShade']:
                                        if attr in color_elem.attrib:
                                            del color_elem.attrib[attr]
                            parent.insert(hyperlink_index + i, child)
                        parent.remove(hyperlink)
                        hyperlinks_removed += 1

                # Replace text
                for text_elem in root.findall('.//w:t', namespaces):
                    if text_elem.text:
                        original_text = text_elem.text
                        new_text = self.replace_keywords_in_text(original_text)
                        if new_text != original_text:
                            text_elem.text = new_text
                            text_replacements += 1

                tree.write(footer_file, encoding='utf-8', xml_declaration=True)

            # Recreate the DOCX file
            print("Rebuilding DOCX file...")
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zip_out:
                for file_path in temp_dir.rglob('*'):
                    if file_path.is_file():
                        # Calculate the path within the zip
                        arc_path = file_path.relative_to(temp_dir)
                        zip_out.write(file_path, arc_path)

            print(f"✓ Removed {images_removed} images/objects")
            print(f"✓ Removed {hyperlinks_removed} hyperlinks (text preserved)")
            print(f"✓ Made {text_replacements} text replacements")

        return output_path

    def remove_selected_images_simple(self, docx_path):
        """
        Simple, aggressive SELECTIVE image removal - runs FIRST before any other processing.
        Only removes images that match the hashes in self.image_hashes_to_remove.
        """
        print("=" * 70)
        print("STEP 0: Simple Selective Image Removal (Before Processing)")
        print("=" * 70)

        # If no selection, remove all images
        if not self.image_hashes_to_remove:
            print("  No image selection - will remove ALL images")
            remove_all = True
        else:
            print(f"  Will remove {len(self.image_hashes_to_remove)} selected images")
            remove_all = False

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_dir = Path(temp_dir)

            # Extract DOCX
            print("  Extracting DOCX...")
            with zipfile.ZipFile(docx_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)

            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            }

            ET.register_namespace('w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')
            ET.register_namespace('r', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships')
            ET.register_namespace('', 'http://schemas.openxmlformats.org/package/2006/relationships')

            # 1. Build map of media files to their hashes
            print("\n1. Analyzing media files...")
            media_dir = temp_dir / 'word' / 'media'
            images_to_remove = set()  # filenames to remove

            if media_dir.exists():
                for media_file in media_dir.iterdir():
                    if media_file.is_file():
                        try:
                            with open(media_file, 'rb') as f:
                                image_data = f.read()
                                image_hash = self.calculate_image_hash(image_data)

                                # Check if this image should be removed
                                if remove_all or image_hash in self.image_hashes_to_remove:
                                    images_to_remove.add(media_file.name)
                                    print(f"    ✓ Marked for removal: {media_file.name}")
                                else:
                                    print(f"    ○ Keeping: {media_file.name}")
                        except Exception as e:
                            print(f"    ✗ Error analyzing {media_file.name}: {e}")

            print(f"  Total images to remove: {len(images_to_remove)}")

            if len(images_to_remove) == 0:
                print("  No images to remove - returning original file")
                return docx_path

            # 2. Find relationship IDs for images to remove
            print("\n2. Finding relationship IDs...")
            rel_ids_to_remove = set()

            rels_dir = temp_dir / 'word' / '_rels'
            if rels_dir.exists():
                for rels_file in rels_dir.glob('*.xml.rels'):
                    try:
                        tree = ET.parse(rels_file)
                        root = tree.getroot()

                        for rel in root.findall(
                                './/{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
                            target = rel.get('Target', '')
                            if 'media/' in target:
                                filename = Path(target).name
                                if filename in images_to_remove:
                                    rel_id = rel.get('Id')
                                    rel_ids_to_remove.add(rel_id)
                                    print(f"    Found: {filename} -> {rel_id}")
                    except Exception as e:
                        print(f"    Warning: Error reading {rels_file.name}: {e}")

            print(f"  Total relationship IDs to remove: {len(rel_ids_to_remove)}")

            # 3. Remove image runs from XML files
            print("\n3. Removing image runs from documents...")

            def remove_selected_image_runs(xml_file):
                if not xml_file.exists():
                    return 0

                try:
                    tree = ET.parse(xml_file)
                    root = tree.getroot()
                    parent_map = {c: p for p in tree.iter() for c in p}

                    runs_removed = 0
                    runs_to_remove = []

                    # Find runs that contain images we want to remove
                    for run in root.findall('.//w:r', namespaces):
                        should_remove = False

                        # Check if this run contains an image
                        for child in run.iter():
                            tag = str(child.tag).lower()
                            if any(img_tag in tag for img_tag in ['drawing', 'pict', 'object', 'picture']):
                                # Check if this specific image should be removed
                                # Look for the relationship ID
                                for blip in child.iter():
                                    if 'blip' in str(blip.tag).lower():
                                        embed_attr = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                                        if embed_attr in blip.attrib:
                                            rel_id = blip.attrib[embed_attr]
                                            if rel_id in rel_ids_to_remove:
                                                should_remove = True
                                                break

                                # If removing all, or if we found a matching rel_id
                                if should_remove:
                                    break

                        if should_remove:
                            runs_to_remove.append(run)

                    # Remove the runs
                    for run in runs_to_remove:
                        parent = parent_map.get(run)
                        if parent is not None:
                            parent.remove(run)
                            runs_removed += 1

                    if runs_removed > 0:
                        tree.write(xml_file, encoding='utf-8', xml_declaration=True)
                        print(f"    ✓ Removed {runs_removed} image runs from {xml_file.name}")

                    return runs_removed
                except Exception as e:
                    print(f"    ✗ Error processing {xml_file.name}: {e}")
                    return 0

            total_runs_removed = 0

            # Process main document
            doc_xml = temp_dir / 'word' / 'document.xml'
            total_runs_removed += remove_selected_image_runs(doc_xml)

            # Process headers
            for header_file in (temp_dir / 'word').glob('header*.xml'):
                total_runs_removed += remove_selected_image_runs(header_file)

            # Process footers
            for footer_file in (temp_dir / 'word').glob('footer*.xml'):
                total_runs_removed += remove_selected_image_runs(footer_file)

            # 4. Clean relationships
            print("\n4. Cleaning relationships...")
            rels_cleaned = 0

            if rels_dir.exists():
                for rels_file in rels_dir.glob('*.xml.rels'):
                    try:
                        tree = ET.parse(rels_file)
                        root = tree.getroot()

                        rels_to_remove_list = []
                        for rel in root.findall(
                                './/{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
                            rel_id = rel.get('Id')
                            if rel_id in rel_ids_to_remove:
                                rels_to_remove_list.append(rel)

                        for rel in rels_to_remove_list:
                            root.remove(rel)
                            rels_cleaned += 1

                        if rels_to_remove_list:
                            tree.write(rels_file, encoding='utf-8', xml_declaration=True)
                    except Exception as e:
                        print(f"    Warning: Error cleaning {rels_file.name}: {e}")

            print(f"    ✓ Cleaned {rels_cleaned} relationships")

            # 5. Delete physical image files
            print("\n5. Deleting physical media files...")
            files_deleted = 0

            if media_dir.exists():
                for filename in images_to_remove:
                    image_file = media_dir / filename
                    if image_file.exists():
                        try:
                            image_file.unlink()
                            files_deleted += 1
                        except Exception as e:
                            print(f"    Warning: Could not delete {filename}: {e}")

            print(f"    ✓ Deleted {files_deleted} media files")

            # 6. Rebuild DOCX
            print("\n6. Rebuilding DOCX...")
            output_temp = docx_path.parent / f"{docx_path.stem}_temp_no_selected_images.docx"

            with zipfile.ZipFile(output_temp, 'w', zipfile.ZIP_DEFLATED) as zip_out:
                for file_path in temp_dir.rglob('*'):
                    if file_path.is_file():
                        arc_path = file_path.relative_to(temp_dir)
                        zip_out.write(file_path, arc_path)

            print(f"\n{'=' * 70}")
            print("✅ SELECTIVE IMAGE REMOVAL COMPLETE")
            print(f"{'=' * 70}")
            print(f"✓ Removed {total_runs_removed} image runs")
            print(f"✓ Cleaned {rels_cleaned} relationships")
            print(f"✓ Deleted {files_deleted} media files")
            print()

            return output_temp

    def process_docx_selective(self, input_path, output_path):
        """
        Process DOCX with selective image removal AND complete formatting standardization.
        Includes processing of headers/footers for image removal.
        """
        print("=" * 70)
        print("SELECTIVE MODE - Image Selection + Complete Formatting Cleanup")
        print("=" * 70)
        print()

        if self.image_hashes_to_remove:
            print(f"Will remove {len(self.image_hashes_to_remove)} selected images")
        else:
            print("Will remove ALL images (no selection provided)")
        print()

        # PHASE 1: XML-level preprocessing
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_dir = Path(temp_dir)

            print("Phase 1: Pre-processing at XML level...")
            with zipfile.ZipFile(input_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)

            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
                'rel': 'http://schemas.openxmlformats.org/package/2006/relationships',
                'v': 'urn:schemas-microsoft-com:vml',
                'o': 'urn:schemas-microsoft-com:office:office'
            }

            for prefix, uri in namespaces.items():
                ET.register_namespace(prefix, uri)

            # Build image hash map for selective removal
            print("  Analyzing images...")
            media_dir = temp_dir / 'word' / 'media'
            image_hash_map = {}
            images_to_remove = set()

            if media_dir.exists():
                for media_file in media_dir.iterdir():
                    if media_file.is_file():
                        try:
                            with open(media_file, 'rb') as f:
                                image_data = f.read()
                                image_hash = self.calculate_image_hash(image_data)
                                image_hash_map[media_file.name] = image_hash

                                if self.should_remove_image(image_data):
                                    images_to_remove.add(media_file.name)
                                    print(f"    ✓ Marked for removal: {media_file.name}")
                                else:
                                    print(f"    ○ Keeping: {media_file.name}")
                        except Exception as e:
                            print(f"    ✗ Error: {e}")

            print(f"  Total images to remove: {len(images_to_remove)}\n")

            # Get ALL relationship IDs for images to remove (from all rels files)
            print("  Processing relationships...")
            rel_ids_to_remove = set()
            rels_files_to_update = []

            # Process main document relationships
            rels_file = temp_dir / 'word' / '_rels' / 'document.xml.rels'
            if rels_file.exists():
                rels_files_to_update.append(rels_file)
                tree = ET.parse(rels_file)
                root = tree.getroot()
                for rel in root.findall(
                        './/{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
                    target = rel.get('Target')
                    if target and 'media/' in target:
                        filename = Path(target).name
                        if filename in images_to_remove:
                            rel_id = rel.get('Id')
                            rel_ids_to_remove.add(rel_id)
                            print(f"    Found in document.xml.rels: {filename} -> {rel_id}")

            # Process header/footer relationships
            rels_dir = temp_dir / 'word' / '_rels'
            if rels_dir.exists():
                for rels_file in rels_dir.glob('*.xml.rels'):
                    if rels_file.name == 'document.xml.rels':
                        continue  # Already processed

                    try:
                        tree = ET.parse(rels_file)
                        root = tree.getroot()
                        has_images = False
                        for rel in root.findall(
                                './/{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
                            target = rel.get('Target')
                            if target and 'media/' in target:
                                filename = Path(target).name
                                if filename in images_to_remove:
                                    rel_id = rel.get('Id')
                                    rel_ids_to_remove.add(rel_id)
                                    has_images = True
                                    print(f"    Found in {rels_file.name}: {filename} -> {rel_id}")

                        if has_images:
                            rels_files_to_update.append(rels_file)
                    except Exception as e:
                        print(f"    Warning: Could not process {rels_file.name}: {e}")

            print(f"  Total relationship IDs to remove: {len(rel_ids_to_remove)}\n")

            # REMOVE RELATIONSHIPS FROM .RELS FILES
            print("  Removing relationships from .rels files...")
            for rels_file_path in rels_files_to_update:
                try:
                    tree = ET.parse(rels_file_path)
                    root = tree.getroot()
                    rels_removed = 0

                    # Find and remove relationships
                    for rel in list(root.findall(
                            './/{http://schemas.openxmlformats.org/package/2006/relationships}Relationship')):
                        rel_id = rel.get('Id')
                        if rel_id in rel_ids_to_remove:
                            root.remove(rel)
                            rels_removed += 1

                    if rels_removed > 0:
                        tree.write(rels_file_path, encoding='utf-8', xml_declaration=True)
                        print(f"    ✓ Removed {rels_removed} relationships from {rels_file_path.name}")

                except Exception as e:
                    print(f"    ✗ Error updating {rels_file_path.name}: {e}")

            # Helper function to remove images from a tree (does NOT save)
            def remove_images_from_tree(root, parent_map):
                """Remove drawing references AND their parent runs from tree"""
                runs_to_remove = []
                drawings_found = 0

                # Find all runs (w:r elements)
                for run in root.findall('.//w:r', namespaces):
                    should_remove_run = False

                    # Check for drawings (modern format)
                    for drawing in run.findall('.//w:drawing', namespaces):
                        for blip in drawing.iter():
                            if 'blip' in str(blip.tag).lower():
                                embed_attr = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                                if embed_attr in blip.attrib:
                                    rel_id = blip.attrib[embed_attr]
                                    if rel_id in rel_ids_to_remove:
                                        should_remove_run = True
                                        drawings_found += 1
                                        break
                        if should_remove_run:
                            break

                    # Check for pictures (older format - w:pict)
                    if not should_remove_run:
                        for pict in run.findall('.//w:pict', namespaces):
                            for elem in pict.iter():
                                if 'imagedata' in str(elem.tag).lower():
                                    rel_id = elem.get(
                                        '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                                    if rel_id and rel_id in rel_ids_to_remove:
                                        should_remove_run = True
                                        drawings_found += 1
                                        break
                            if should_remove_run:
                                break

                    # Check for objects (embedded objects)
                    if not should_remove_run:
                        for obj in run.findall('.//w:object', namespaces):
                            for elem in obj.iter():
                                if 'imagedata' in str(elem.tag).lower():
                                    rel_id = elem.get(
                                        '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                                    if rel_id and rel_id in rel_ids_to_remove:
                                        should_remove_run = True
                                        drawings_found += 1
                                        break
                            if should_remove_run:
                                break

                    if should_remove_run:
                        runs_to_remove.append(run)

                # Remove the runs
                runs_removed = 0
                for run in runs_to_remove:
                    parent = parent_map.get(run)
                    if parent is not None:
                        try:
                            parent.remove(run)
                            runs_removed += 1
                        except Exception as e:
                            print(f"      Warning: Could not remove run: {e}")

                return runs_removed, drawings_found

            # Process document.xml - COMPREHENSIVE CLEANUP
            document_xml = temp_dir / 'word' / 'document.xml'
            sdts_removed = 0
            styles_removed = 0
            colors_forced = 0
            shading_removed = 0
            borders_removed = 0
            images_removed = 0
            text_replacements = 0

            if document_xml.exists():
                print("  Processing document.xml...")
                tree = ET.parse(document_xml)
                root = tree.getroot()
                parent_map = {c: p for p in tree.iter() for c in p}

                # 1. REMOVE CONTENT CONTROLS (blue boxes)
                for sdt in root.findall('.//w:sdt', namespaces):
                    parent = parent_map.get(sdt)
                    if parent is not None:
                        sdt_index = list(parent).index(sdt)
                        sdt_content = sdt.find('.//w:sdtContent', namespaces)
                        if sdt_content is not None:
                            for child in list(sdt_content):
                                parent.insert(sdt_index, child)
                                sdt_index += 1
                        parent.remove(sdt)
                        sdts_removed += 1

                # 2. REMOVE ALL PARAGRAPH STYLES (force to Normal)
                for para in root.findall('.//w:p', namespaces):
                    for pPr in para.findall('.//w:pPr', namespaces):
                        for pStyle in pPr.findall('.//w:pStyle', namespaces):
                            pPr.remove(pStyle)
                            styles_removed += 1

                # 3. FORCE ALL TEXT TO BLACK
                for rPr in root.findall('.//w:rPr', namespaces):
                    # Remove existing color elements
                    for color_elem in list(rPr):
                        if 'color' in str(color_elem.tag).lower() or 'highlight' in str(color_elem.tag).lower():
                            rPr.remove(color_elem)

                    # Add black color
                    color_elem = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
                    color_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '000000')

                    # Remove theme attributes
                    for attr in ['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeColor',
                                 '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeTint',
                                 '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeShade']:
                        if attr in color_elem.attrib:
                            del color_elem.attrib[attr]

                    rPr.insert(0, color_elem)
                    colors_forced += 1

                # 4. REMOVE ALL SHADING
                for shd in root.findall('.//w:shd', namespaces):
                    parent = parent_map.get(shd)
                    if parent is not None:
                        parent.remove(shd)
                        shading_removed += 1

                # 5. REMOVE ALL BORDERS
                for pBdr in root.findall('.//w:pBdr', namespaces):
                    parent = parent_map.get(pBdr)
                    if parent is not None:
                        parent.remove(pBdr)
                        borders_removed += 1

                # 6. REMOVE SELECTED IMAGES from document (using same tree, does NOT save yet)
                # Need to rebuild parent map after all the removals above
                parent_map = {c: p for p in tree.iter() for c in p}
                runs_removed, drawings_found = remove_images_from_tree(root, parent_map)
                images_removed = runs_removed
                print(f"    ✓ Removed {runs_removed} runs containing {drawings_found} images from document.xml")

                # 7. REPLACE KEYWORDS
                for text_elem in root.findall('.//w:t', namespaces):
                    if text_elem.text:
                        original = text_elem.text
                        new_text = self.replace_keywords_in_text(original)
                        if new_text != original:
                            text_elem.text = new_text
                            text_replacements += 1

                # NOW save everything at once
                tree.write(document_xml, encoding='utf-8', xml_declaration=True)

                print(f"    ✓ Removed {sdts_removed} content controls")
                print(f"    ✓ Removed {styles_removed} paragraph styles")
                print(f"    ✓ Forced {colors_forced} text runs to black")
                print(f"    ✓ Removed {shading_removed} shading elements")
                print(f"    ✓ Removed {borders_removed} borders")
                print(f"    ✓ Made {text_replacements} keyword replacements")

            # Helper function for header/footer files
            def process_header_footer_file(xml_path, xml_name):
                """Process header/footer XML file - remove images and clean formatting"""
                if not xml_path.exists():
                    return 0

                try:
                    tree = ET.parse(xml_path)
                    root = tree.getroot()
                    parent_map = {c: p for p in tree.iter() for c in p}

                    # Remove images
                    runs_removed, drawings_found = remove_images_from_tree(root, parent_map)

                    # Force all text to black (same as document)
                    colors_forced = 0
                    for rPr in root.findall('.//w:rPr', namespaces):
                        # Remove existing color elements
                        for color_elem in list(rPr):
                            if 'color' in str(color_elem.tag).lower() or 'highlight' in str(color_elem.tag).lower():
                                rPr.remove(color_elem)

                        # Add black color
                        color_elem = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
                        color_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '000000')

                        for attr in ['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeColor',
                                     '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeTint',
                                     '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}themeShade']:
                            if attr in color_elem.attrib:
                                del color_elem.attrib[attr]

                        rPr.insert(0, color_elem)
                        colors_forced += 1

                    # Remove paragraph styles
                    styles_removed = 0
                    parent_map = {c: p for p in tree.iter() for c in p}
                    for para in root.findall('.//w:p', namespaces):
                        for pPr in para.findall('.//w:pPr', namespaces):
                            for pStyle in pPr.findall('.//w:pStyle', namespaces):
                                pPr.remove(pStyle)
                                styles_removed += 1

                    # Remove shading
                    shading_removed = 0
                    parent_map = {c: p for p in tree.iter() for c in p}
                    for shd in root.findall('.//w:shd', namespaces):
                        parent = parent_map.get(shd)
                        if parent is not None:
                            parent.remove(shd)
                            shading_removed += 1

                    if runs_removed > 0 or colors_forced > 0 or styles_removed > 0 or shading_removed > 0:
                        tree.write(xml_path, encoding='utf-8', xml_declaration=True)
                        print(
                            f"    ✓ Processed {xml_name}: {runs_removed} image runs, {colors_forced} colors, {styles_removed} styles, {shading_removed} shading")

                    return runs_removed
                except Exception as e:
                    print(f"    ✗ Error processing {xml_name}: {e}")
                    return 0

            # Process ALL header and footer files
            print("  Processing headers and footers...")
            word_dir = temp_dir / 'word'
            header_footer_images = 0

            if word_dir.exists():
                # Process all header files
                for header_file in word_dir.glob('header*.xml'):
                    count = process_header_footer_file(header_file, header_file.name)
                    header_footer_images += count

                # Process all footer files
                for footer_file in word_dir.glob('footer*.xml'):
                    count = process_header_footer_file(footer_file, footer_file.name)
                    header_footer_images += count

            images_removed += header_footer_images

            # Neutralize theme files
            print("  Neutralizing theme files...")
            theme_dir = temp_dir / 'word' / 'theme'
            if theme_dir.exists():
                for theme_file in theme_dir.glob('*.xml'):
                    try:
                        tree = ET.parse(theme_file)
                        root = tree.getroot()
                        for element in root.iter():
                            if 'srgbClr' in str(element.tag):
                                element.set('val', '000000')
                            elif 'sysClr' in str(element.tag):
                                element.set('val', 'windowText')
                                element.set('lastClr', '000000')
                        tree.write(theme_file, encoding='utf-8', xml_declaration=True)
                    except:
                        pass

            # Neutralize styles.xml
            print("  Neutralizing styles.xml...")
            styles_xml = temp_dir / 'word' / 'styles.xml'
            if styles_xml.exists():
                tree = ET.parse(styles_xml)
                root = tree.getroot()
                parent_map = {c: p for p in tree.iter() for c in p}

                elements_to_remove = []
                for elem in root.iter():
                    if 'color' in str(elem.tag).lower() or 'shd' in str(elem.tag).lower():
                        elements_to_remove.append(elem)

                for elem in elements_to_remove:
                    parent = parent_map.get(elem)
                    if parent is not None:
                        try:
                            parent.remove(elem)
                        except:
                            pass

                tree.write(styles_xml, encoding='utf-8', xml_declaration=True)

            # Save to temp file WITHOUT deleting images yet
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_preprocessed:
                preprocessed_path = tmp_preprocessed.name

            with zipfile.ZipFile(preprocessed_path, 'w', zipfile.ZIP_DEFLATED) as zip_out:
                for file_path in temp_dir.rglob('*'):
                    if file_path.is_file():
                        arc_path = file_path.relative_to(temp_dir)
                        zip_out.write(file_path, arc_path)

        # PHASE 2: Python-docx processing for remaining cleanup
        print("\nPhase 2: Processing with python-docx for final cleanup...")

        if not DOCX_AVAILABLE:
            print("  python-docx not available, skipping to final cleanup...")
            temp_output = preprocessed_path
        else:
            try:
                doc = Document(preprocessed_path)

                # Process all paragraphs
                for para in doc.paragraphs:
                    try:
                        if para.style.name != 'Normal':
                            para.style = doc.styles['Normal']
                    except:
                        pass

                    for run in para.runs:
                        self.standardize_run_formatting(run)

                # Process tables
                for table in doc.tables:
                    for row in table.rows:
                        self.remove_table_row_shading(row)
                        for cell in row.cells:
                            self.remove_table_cell_shading(cell)
                            for paragraph in cell.paragraphs:
                                try:
                                    if paragraph.style.name != 'Normal':
                                        paragraph.style = doc.styles['Normal']
                                except:
                                    pass
                                for run in paragraph.runs:
                                    self.standardize_run_formatting(run)

                # Process headers/footers
                for section in doc.sections:
                    if section.header:
                        for paragraph in section.header.paragraphs:
                            try:
                                if paragraph.style.name != 'Normal':
                                    paragraph.style = doc.styles['Normal']
                            except:
                                pass
                            for run in paragraph.runs:
                                self.standardize_run_formatting(run)

                    if section.footer:
                        for paragraph in section.footer.paragraphs:
                            try:
                                if paragraph.style.name != 'Normal':
                                    paragraph.style = doc.styles['Normal']
                            except:
                                pass
                            for run in paragraph.runs:
                                self.standardize_run_formatting(run)

                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_phase2:
                    temp_output = tmp_phase2.name

                doc.save(temp_output)

                try:
                    os.unlink(preprocessed_path)
                except:
                    pass

            except Exception as e:
                print(f"  Warning: python-docx processing failed: {e}")
                print("  Continuing with phase 1 output...")
                temp_output = preprocessed_path

        # PHASE 3: Final cleanup - NOW remove the physical image files
        print("\nPhase 3: Final cleanup - removing physical image files...")

        with tempfile.TemporaryDirectory() as final_temp_dir:
            final_temp_dir = Path(final_temp_dir)

            with zipfile.ZipFile(temp_output, 'r') as zip_ref:
                zip_ref.extractall(final_temp_dir)

            # Delete the physical image files
            media_dir = final_temp_dir / 'word' / 'media'
            files_deleted = 0
            if images_to_remove and media_dir.exists():
                for filename in images_to_remove:
                    image_file = media_dir / filename
                    if image_file.exists():
                        try:
                            image_file.unlink()
                            files_deleted += 1
                        except Exception as e:
                            print(f"  Warning: Could not delete {filename}: {e}")

            print(f"  Removed {files_deleted} physical image files")

            # Rebuild final DOCX
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zip_out:
                for file_path in final_temp_dir.rglob('*'):
                    if file_path.is_file():
                        arc_path = file_path.relative_to(final_temp_dir)
                        zip_out.write(file_path, arc_path)

        # Clean up temp file
        try:
            os.unlink(temp_output)
        except:
            pass

        print(f"\n{'=' * 70}")
        print("✅ PROCESSING COMPLETE")
        print(f"{'=' * 70}")
        print(
            f"✓ Removed {images_removed} runs containing images (kept {len(image_hash_map) - len(images_to_remove)} images)")
        print(f"✓ Removed {sdts_removed} content controls")
        print(f"✓ Reset {styles_removed} paragraph styles to Normal")
        print(f"✓ Forced {colors_forced} text runs to black")
        print(f"✓ Removed {shading_removed} shading elements")
        print(f"✓ Removed {borders_removed} borders")
        print(f"✓ Made {text_replacements} keyword replacements")
        print(f"✓ Neutralized themes and styles")
        print(f"✓ Removed {len(rel_ids_to_remove)} image relationships")
        print()

        return output_path
    def process_html_file(self, input_path, output_path):
        """Process HTML file - remove images and replace keywords, keep structure and CSS"""
        if not HTML_AVAILABLE:
            raise ImportError("beautifulsoup4 not installed. Run: pip install beautifulsoup4")

        with open(input_path, 'r', encoding='utf-8') as file:
            content = file.read()

        soup = BeautifulSoup(content, 'html.parser')

        # Remove all image-related elements
        images_removed = 0
        for img_tag in soup.find_all(['img', 'picture', 'svg', 'canvas']):
            img_tag.decompose()
            images_removed += 1

        # Remove background images from CSS (inline styles)
        for element in soup.find_all(style=True):
            style = element.get('style', '')
            if 'background-image' in style.lower():
                style = re.sub(r'background-image\s*:[^;]*;?', '', style, flags=re.IGNORECASE)
                element['style'] = style
                images_removed += 1

        # Remove background images from CSS in style tags
        for style_tag in soup.find_all('style'):
            if style_tag.string:
                css_content = style_tag.string
                if 'background-image' in css_content.lower():
                    css_content = re.sub(r'background-image\s*:[^;]*;?', '', css_content, flags=re.IGNORECASE)
                    style_tag.string = css_content
                    images_removed += 1

        # Remove hyperlinks but keep text content
        hyperlinks_removed = 0
        for link in soup.find_all('a'):
            # Get the text content
            link_text = link.get_text()
            # Replace the link with just the text
            link.replace_with(link_text)
            hyperlinks_removed += 1

        # Replace keywords in text nodes
        text_replacements = 0

        def replace_text_nodes(element):
            nonlocal text_replacements
            if element.string:
                original_text = str(element.string)
                new_text = self.replace_keywords_in_text(original_text)
                if new_text != original_text:
                    element.string.replace_with(new_text)
                    text_replacements += 1
            else:
                for child in list(element.children):
                    if hasattr(child, 'name') and child.name:  # It's a tag
                        replace_text_nodes(child)
                    elif hasattr(child, 'replace_with'):  # It's a text node
                        original_text = str(child)
                        new_text = self.replace_keywords_in_text(original_text)
                        if new_text != original_text:
                            child.replace_with(new_text)
                            text_replacements += 1

        if soup.body:
            replace_text_nodes(soup.body)

        # Write the processed HTML
        with open(output_path, 'w', encoding='utf-8') as file:
            file.write(str(soup))

        print(f"✓ Removed {images_removed} images/graphics")
        print(f"✓ Removed {hyperlinks_removed} hyperlinks (text preserved)")
        print(f"✓ Made {text_replacements} text replacements")

        return output_path

    def process_txt_file(self, input_path, output_path):
        """Process plain text file - replace keywords only"""
        try:
            with open(input_path, 'r', encoding='utf-8') as file:
                content = file.read()
        except UnicodeDecodeError:
            with open(input_path, 'r', encoding='latin-1') as file:
                content = file.read()

        # Replace keywords
        processed_content = self.replace_keywords_in_text(content)

        # Count replacements
        text_replacements = 0
        for original, replacement in self.keyword_replacements.items():
            if original.startswith(r'\b') or original.startswith(r'['):
                text_replacements += len(re.findall(original, content, flags=re.IGNORECASE))
            else:
                text_replacements += len(re.findall(re.escape(original), content, flags=re.IGNORECASE))

        # Write processed content
        with open(output_path, 'w', encoding='utf-8') as file:
            file.write(processed_content)

        print(f"✓ Made {text_replacements} text replacements")

        return output_path

    def blind_file(self, input_path, output_path=None, method='safe'):
        """
        Main function to blind a file - removes images and replaces keywords

        Args:
            input_path (str): Path to input file
            output_path (str): Path to output file (optional)
            method (str): Processing method - 'safe' (python-docx) or 'xml' (direct XML)
        """
        input_path = Path(input_path)

        if not input_path.exists():
            raise FileNotFoundError(f"Input file not found: {input_path}")

        # Generate output path if not provided
        if output_path is None:
            output_path = input_path.parent / f"{input_path.stem}_blinded{input_path.suffix}"
        else:
            output_path = Path(output_path)

        # Ensure output directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)

        extension = input_path.suffix.lower()

        try:
            print(f"Processing file: {input_path}")
            print(f"Output will be saved to: {output_path}")
            print()

            if extension == '.docx':
                # Use selective method if image hashes are provided
                if self.image_hashes_to_remove:
                    result = self.process_docx_selective(input_path, output_path)
                    print("✓ DOCX processed with selective image removal")
                elif method == 'xml':
                    result = self.process_docx_xml_safe(input_path, output_path)
                    print("✓ DOCX processed with XML method")
                else:
                    result = self.process_docx_safe(input_path, output_path)
                    print("✓ DOCX processed with safe method")

            elif extension == '.html' or extension == '.htm':
                result = self.process_html_file(input_path, output_path)
                print("✓ HTML processed")

            elif extension == '.txt':
                result = self.process_txt_file(input_path, output_path)
                print("✓ Text file processed")

            else:
                raise ValueError(f"Unsupported file type: {extension}. Supported: .docx, .html, .htm, .txt")

            return result

        except Exception as e:
            print(f"Error processing file: {e}")
            return None


def main():
    """Interactive interface for file blinding"""
    print("=== File Blinding Tool (Structure Preserving) ===")
    print("Removes images and replaces keywords while keeping ALL original formatting")
    print("Supports: DOCX, HTML, TXT")
    print()

    # Get input file path
    while True:
        input_file = input("Enter input file path: ").strip().strip('"\'')
        if not input_file:
            print("Please enter a file path.")
            continue

        input_path = Path(input_file)
        if not input_path.exists():
            print(f"Error: File not found at '{input_file}'")
            continue

        # Check if file type is supported
        supported_extensions = ['.docx', '.html', '.htm', '.txt']
        if input_path.suffix.lower() not in supported_extensions:
            print(f"Error: Unsupported file type '{input_path.suffix}'")
            print(f"Supported formats: {', '.join(supported_extensions)}")
            continue

        break

    # Get output file path
    print()
    output_file = input("Enter output file path (or press Enter for auto-generated): ").strip().strip('"\'')

    if not output_file:
        output_file = None
        print("Auto-generating output filename...")

    # For DOCX files, ask about processing method
    method = 'safe'
    if Path(input_file).suffix.lower() == '.docx':
        print()
        print("DOCX Processing methods:")
        print("1. Safe (recommended) - Uses python-docx library, very reliable")
        print("2. XML - Direct XML processing, maximum structure preservation")

        method_choice = input("Choose processing method (1-2, default=1): ").strip()
        method = 'xml' if method_choice == '2' else 'safe'

    # Keyword replacements
    replacements = {
        # Sensitive terms
        "confidential": "[REDACTED]",
        "secret": "[REDACTED]",
        "internal": "[INTERNAL]",
        "proprietary": "[PROPRIETARY]",
        "classified": "[CLASSIFIED]",

        # Personal info patterns (regex)
        r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b': '[EMAIL]',
        r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b': '[PHONE]',
        r'\b\d{3}-\d{2}-\d{4}\b': '[SSN]',
        r'\b\d{1,2}/\d{1,2}/\d{4}\b': '[DATE]',
        r'\$[\d,]+\.?\d*': '[AMOUNT]',

        # Add your specific terms here!
        # "Your Company Name": "[COMPANY]",
        # "John Smith": "[PERSON]",
    }

    print()
    print("Formatting standardization:")
    print("✓ Font will be changed to Calibri")
    print("✓ All text colors will be changed to black")
    print("✓ All highlighting and shading will be removed")
    print("✓ All table cell backgrounds will be removed")
    print("✓ All borders will be removed")
    print("✓ All hyperlinks will be removed (text preserved)")
    print("✓ Document themes will be removed")
    print()

    blinder = FileBlinder(
        keyword_replacements=replacements,
        standardize_formatting=True,
        font_name="Calibri",
        font_size=11,  # You can change this if needed
        font_color_black=True,
        grey_shading=False  # Changed to False to remove all highlighting
    )

    print("Processing file...")
    print()

    try:
        result = blinder.blind_file(input_file, output_file, method)

        if result:
            print()
            print("🎉 SUCCESS!")
            print(f"Blinded file created at: {result}")
            print()
            print("✓ All formatting and structure preserved")
            print("✓ Font standardized to Calibri")
            print("✓ All text changed to black with no highlighting")
            print("✓ All table cell backgrounds removed")
            print("✓ All borders and shading removed")
            print("✓ All hyperlinks removed (text preserved)")
            print("✓ Document themes removed")
            print("✓ Tables, spacing, layout maintained")
            print("✓ Images and graphics removed")
            print("✓ Keywords replaced as configured")
        else:
            print("✗ Failed to process file.")

    except Exception as e:
        print(f"✗ Error: {e}")

    print()
    input("Press Enter to exit...")


if __name__ == "__main__":
    main()